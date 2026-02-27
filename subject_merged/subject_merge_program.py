import os
import re
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from datetime import datetime, timedelta

# 현재 시간에서 하루 전 빼기
yesterday = datetime.now() - timedelta(days=1)

date_str = yesterday.strftime("%Y%m%d") 
weekday_str = yesterday.strftime("%A")    

weekday_dict = {
    "Monday": "월요일",
    "Tuesday": "화요일",
    "Wednesday": "수요일",
    "Thursday": "목요일",
    "Friday": "금요일",
    "Saturday": "토요일",
    "Sunday": "일요일"
}

weekday_kor = weekday_dict.get(weekday_str, weekday_str)

filename = f"{date_str}_{weekday_kor}_피드백 통합본.docx"

# 과목 순서
SUBJECT_ORDER = ["Grammar", "Reading", "Writing", "Listening", "Voca"]

# 출결 우선순위 정리
def update_attendance(current, new):
    priority = {'결석': 3, '지각': 2, 'O': 1}
    def extract_status(text):
        for key in priority:
            if key in text:
                return key
        return None

    cur_status = extract_status(current) if current else None
    new_status = extract_status(new)

    if new_status and (not cur_status or priority[new_status] > priority.get(cur_status, 0)):
        return new
    return current

# 특수기호 제거 함수
def clean_text(text):
    return re.sub(r'[^\w가-힣()\[\]{}: %/+★☆◉▒\s]', '', text).strip()

# 맑은 고딕 폰트
def set_font_korean(paragraph):
    for run in paragraph.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# docx 파일 1개
def parse_docx_file(file_path, student_data):
    doc = Document(file_path)
    full_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    blocks = re.split(r'(『[^』]+학생』)', full_text)
    if len(blocks) < 3:
        print(f"[경고] '{os.path.basename(file_path)}'에서 학생 블[록이 제대로 추출되지 않음.")
        return

    for i in range(1, len(blocks), 2):
        raw_name = blocks[i]
        name = clean_text(raw_name.replace("『", "").replace("』", "").replace("학생", "").strip())
        content = blocks[i + 1]

        # 출결
        att_match = re.search(r': : 출결\([^)]+\): [^\n]+', content)
        if att_match:
            student_data[name]["출결"] = update_attendance(student_data[name]["출결"], att_match.group(0))

        # 과목별 피드백
        subject_iter = list(re.finditer(r'(▶[^\n]+)', content))
        for j, m in enumerate(subject_iter):
            subject_line = m.group(1)
            subject_name = None
            for s in SUBJECT_ORDER:
                if s.lower() in subject_line.lower():
                    subject_name = s
                    break
            if subject_name:
                start = m.start()
                end = subject_iter[j + 1].start() if j + 1 < len(subject_iter) else len(content)
                section_text = content[start:end].strip()
                if section_text not in student_data[name]["과목"][subject_name]:
                    student_data[name]["과목"][subject_name].append(section_text)

# 리포트 생성
def merge_feedbacks(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    student_data = defaultdict(lambda: {"출결": None, "과목": defaultdict(list)})

    # 모든 파일 처리
    for file in os.listdir(input_folder):
        if file.endswith(".docx"):
            parse_docx_file(os.path.join(input_folder, file), student_data)

    # 학생별 리포트 생성
    for name, info in student_data.items():
        doc = Document()

        # 페이지 여백 좁게 설정
        section = doc.sections[0]
        from docx.shared import Inches
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # 이름 줄
        p = doc.add_paragraph()
        run = p.add_run(f'『{name} 학생』')
        run.bold = True
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # 출결
        if info["출결"]:
            para = doc.add_paragraph(info["출결"])
            set_font_korean(para)

        # 과목별 피드백
        for subject in SUBJECT_ORDER:
            if subject in info["과목"] and info["과목"][subject]:
                if subject == "Listening":
                    p_blank = doc.add_paragraph("")
                    set_font_korean(p_blank)
                for section in info["과목"][subject]:
                    para = doc.add_paragraph(section)
                    set_font_korean(para)
                # 과목 블록 사이에 한 줄 띄우기
                doc.add_paragraph("")

        # 저장
        save_path = os.path.join(output_folder, f"{name}_리포트.docx")
        doc.save(save_path)
        print(f"[완료] {name} 리포트 저장됨.")

        combined_doc = Document()

    # 페이지 여백 좁게 설정
    section = combined_doc.sections[0]
    from docx.shared import Inches
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    for name, info in student_data.items():
        p = combined_doc.add_paragraph()
        run = p.add_run(f'『{name} 학생』')
        run.bold = True
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        if info["출결"]:
            para = combined_doc.add_paragraph(info["출결"])
            set_font_korean(para)

        for subject in SUBJECT_ORDER:
            if subject in info["과목"] and info["과목"][subject]:
                if subject == "Listening":
                    p_blank = combined_doc.add_paragraph("")
                    set_font_korean(p_blank)
                for section_text in info["과목"][subject]:
                    para = combined_doc.add_paragraph(section_text)
                    set_font_korean(para)

                combined_doc.add_paragraph("")  # 줄바꿈

    combined_save_path = os.path.join(output_folder, filename)
    combined_doc.save(combined_save_path)
    print(f"[완료] 전체 학생 리포트 저장됨: {combined_save_path}")

# 실행
if __name__ == "__main__":
    input_folder = r"C:\Users\Seo Yeon\OneDrive\바탕 화면\피드백 파일 모음\피드백 통합본 파일"
    output_folder = "피드백 통합본 파일 output"

    merge_feedbacks(input_folder, output_folder)


