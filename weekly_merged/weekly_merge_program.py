from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
import re
from collections import defaultdict
import datetime
import os


def extract_blocks_by_student(doc, day_label):
    blocks = defaultdict(list)
    current_student = None
    current_block = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("#"):
            continue
        
        match = re.match(r'^『(.+?) 학생』', text)
        if match:
            if current_student and current_block:
                blocks[current_student].append((day_label, current_block))
            current_student = match.group(1)
            current_block = [para]
        elif current_student:
            current_block.append(para)

    if current_student and current_block:
        blocks[current_student].append((day_label, current_block))

    return blocks

def get_ordered_student_list(doc1, doc2):
    ordered_students = []
    seen = set()

    for para in doc1.paragraphs + doc2.paragraphs:
        match = re.match(r'^『(.+?) 학생』', para.text.strip())
        if match:
            name = match.group(1)
            if name not in seen:
                ordered_students.append(name)
                seen.add(name)

    return ordered_students

# 자동 주차 생성
def get_custom_week_label(today=None):
    if today is None:
        today = datetime.date.today()

    this_year = today.year
    this_month = today.month

    # 월요일
    current_monday = today - datetime.timedelta(days=today.weekday())

    # 이번 달 1일
    first_day = datetime.date(this_year, this_month, 1)
    first_day_weekday = first_day.weekday()
    first_day_monday = first_day - datetime.timedelta(days=first_day_weekday)

    # 1일이 토/일이면 - 다음 주부터 1주차
    if first_day.weekday() in [5, 6]:
        first_week_monday = first_day_monday + datetime.timedelta(weeks=1)
    else:
        first_week_monday = first_day_monday

    # 현재 주차
    week_number = ((current_monday - first_week_monday).days // 7) + 1

    return f"{this_month}월 {week_number}주차"


def find_docx_files(folder):
    files = [f for f in os.listdir(folder) if f.endswith('.docx') and '피드백 통합본' in f]
    if len(files) != 2:
        raise ValueError(f"폴더 내 '피드백 통합본' 포함된 docx 파일이 정확히 2개여야 합니다. 현재 {len(files)}개 발견됨.")
    return sorted(files)


def extract_weekday_from_filename(filename):
    weekdays = ['월요일', '화요일', '수요일', '목요일', '금요일', '토요일', '일요일']
    for day in weekdays:
        if day in filename:
            return day
    raise ValueError(f"파일명에서 요일을 찾을 수 없습니다: {filename}")


def merge_student_blocks_auto(input_folder, output_folder):
    week_label = get_custom_week_label()
    
    file_names = find_docx_files(input_folder)
    file_paths = [os.path.join(input_folder, f) for f in file_names]

    weekday_labels = [extract_weekday_from_filename(name) for name in file_names]
    weekday_short = ''.join([label[0] for label in weekday_labels])  # 예: 화목

    filename = f"{week_label}_{weekday_short}_피드백 통합본.docx"
    output_file = os.path.join(output_folder, filename)

    doc1 = Document(file_paths[0])
    doc2 = Document(file_paths[1])

    blocks1 = extract_blocks_by_student(doc1, weekday_labels[0])
    blocks2 = extract_blocks_by_student(doc2, weekday_labels[1])
    ordered_students = get_ordered_student_list(doc1, doc2)

    merged_doc = Document()
    section = merged_doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    for student in ordered_students:
        para = merged_doc.add_paragraph()
        run = para.add_run(f'『{student} 학생』')
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        merged_doc.add_paragraph(f'♥{week_label} 피드백♥')
        merged_doc.add_paragraph()

        for _, paras in blocks1.get(student, []):
            for i, p in enumerate(paras):
                if i == 0:
                    continue
                merged_doc.add_paragraph(p.text)
                if p.text.strip().startswith('▶'):
                    merged_doc.add_paragraph()

        for _, paras in blocks2.get(student, []):
            for i, p in enumerate(paras):
                if i == 0:
                    continue
                merged_doc.add_paragraph(p.text)
                if p.text.strip().startswith('▶'):
                    merged_doc.add_paragraph()

        merged_doc.add_paragraph()

    merged_doc.save(output_file)
    print(f"[완료] 저장 파일명: {output_file}")

merge_student_blocks_auto(
    input_folder=r"C:\Users\Seo Yeon\OneDrive\바탕 화면\피드백 파일 모음\요일별 피드백 통합본",
    output_folder=r"C:\Users\Seo Yeon\OneDrive\바탕 화면\피드백 파일 모음\요일별 피드백 통합본 output"
)
